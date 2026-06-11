from docx import Document

def extract_text_from_docx(file_path):
    """
    Extract all text from a DOCX file
    """
    doc = Document(file_path)
    full_text = []
    
    # Extract paragraphs
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    
    return n.join(full_text)

def main():
    # Extract from the main resume file
    try:
        text = extract_text_from_docx("/home/lucas/Documents/job_search/resumes/resume_neo_hong_chuan.docx")
        print("Extracted text from resume:")
        print(text)
    except Exception as e:
        print(f"Error reading main resume: {e}")

if __name__ == "__main__":
    main()

